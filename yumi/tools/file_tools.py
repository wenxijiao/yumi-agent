from __future__ import annotations

import csv
import io
import json
from pathlib import Path

from yumi.logging_config import get_logger

logger = get_logger(__name__)

MAX_CONTENT_CHARS = 50_000
_TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".log",
    ".ini",
    ".cfg",
    ".conf",
    ".yaml",
    ".yml",
    ".toml",
    ".xml",
    ".html",
    ".htm",
    ".css",
    ".js",
    ".ts",
    ".py",
    ".java",
    ".c",
    ".cpp",
    ".h",
    ".hpp",
    ".go",
    ".rs",
    ".rb",
    ".php",
    ".sh",
    ".bash",
    ".zsh",
    ".bat",
    ".ps1",
    ".sql",
    ".r",
    ".m",
    ".swift",
    ".kt",
    ".scala",
    ".lua",
    ".pl",
    ".tex",
    ".env",
    ".gitignore",
    ".dockerignore",
    ".editorconfig",
}
_IMAGE_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".webp",
    ".bmp",
    ".tiff",
    ".tif",
    ".ico",
}


def _truncate(text: str, limit: int = MAX_CONTENT_CHARS) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + f"\n\n... [truncated, showing first {limit} of {len(text)} characters]"


_SENSITIVE_FILE_NAMES = {
    "id_rsa",
    "id_dsa",
    "id_ecdsa",
    "id_ed25519",
    "authorized_keys",
    "known_hosts",
    "credentials",
    "config.json",
}
_SENSITIVE_DIR_NAMES = {".ssh", ".aws", ".gnupg", ".yumi", ".kube"}


def _uploads_root() -> Path:
    return (Path.home() / ".yumi" / "uploads").resolve()


def _is_sensitive_path(path: Path) -> bool:
    # Uploaded files live under ~/.yumi/uploads and are meant to be read back
    # (the upload service restricts what extensions can land there), so allow
    # that subtree even though it sits inside the otherwise-blocked .yumi dir.
    try:
        resolved = path.resolve()
        uploads = _uploads_root()
        if resolved == uploads or uploads in resolved.parents:
            return False
    except (OSError, RuntimeError):
        pass

    parts = {p for p in path.parts}
    if parts & _SENSITIVE_DIR_NAMES:
        return True
    name = path.name.lower()
    if name in _SENSITIVE_FILE_NAMES:
        return True
    if name.endswith(".pem") or name.endswith(".key"):
        return True
    return False


def _safe_resolve(file_path: str) -> Path:
    return Path(file_path).expanduser().resolve()


def _read_text_file(path: Path, encoding: str = "utf-8") -> str:
    try:
        return path.read_text(encoding=encoding)
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def _read_docx(path: Path) -> str:
    try:
        import docx
    except ImportError:
        return "Error: python-docx is missing (it ships with yumi-agent). Reinstall: pip install --force-reinstall yumi-agent"

    document = docx.Document(str(path))
    parts: list[str] = []
    for para in document.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    if not parts:
        return "No extractable text in this Word document (it may be mostly images or empty)."
    return "\n\n".join(parts)


def _read_pdf(path: Path) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return (
            "Error: PyPDF2 is missing (it ships with yumi-agent). Reinstall: pip install --force-reinstall yumi-agent"
        )

    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- Page {i} ---\n{text.strip()}")
    if not pages:
        return "The PDF contains no extractable text (it may be image-based)."
    return "\n\n".join(pages)


def _read_csv_file(path: Path) -> str:
    raw = _read_text_file(path)
    try:
        dialect = csv.Sniffer().sniff(raw[:4096])
    except csv.Error:
        dialect = None

    reader = csv.reader(io.StringIO(raw), dialect=dialect or "excel")
    rows = list(reader)
    if not rows:
        return "The CSV file is empty."

    lines: list[str] = []
    total = len(rows)
    for i, row in enumerate(rows):
        if i == 0:
            lines.append("Header: " + " | ".join(row))
        else:
            lines.append(f"Row {i}: " + " | ".join(row))
        if len(lines) >= 500:
            lines.append(f"... [{total - 500} more rows not shown]")
            break
    return "\n".join(lines)


def _read_json_file(path: Path) -> str:
    raw = _read_text_file(path)
    try:
        data = json.loads(raw)
        return json.dumps(data, ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        return raw


def read_file(file_path: str) -> str:
    path = _safe_resolve(file_path)

    if _is_sensitive_path(path):
        return (
            f"Error: refusing to read sensitive path '{path}'. "
            "Credentials, SSH keys, and Yumi's own config are blocked from the read_file tool."
        )

    if not path.exists():
        return f"Error: File not found: {path}"
    if not path.is_file():
        return f"Error: '{path}' is not a file (it may be a directory)."

    ext = path.suffix.lower()
    size = path.stat().st_size
    header = f"[File: {path.name} | Size: {size:,} bytes | Type: {ext or 'no extension'}]\n\n"

    if ext in _IMAGE_EXTENSIONS:
        info = f"This is an image file ({ext}). Size: {size:,} bytes."
        try:
            from PIL import Image as _PILImage

            with _PILImage.open(path) as img:
                w, h = img.size
                info += f" Dimensions: {w}x{h} pixels. Format: {img.format or ext.upper()}."
        except Exception:
            # PIL missing or the image is unreadable/corrupt — fall back to
            # metadata-only output. Logged at debug so it isn't silent.
            logger.debug("Could not read image metadata for %s", path, exc_info=True)
        info += (
            " If this chat uses a vision-capable model, the image may be embedded in the user message. "
            "If vision is disabled or the model is text-only, you cannot see pixel content—"
            "describe only metadata here or ask the user to switch model or describe the image."
        )
        return header + info

    try:
        if ext == ".pdf":
            content = _read_pdf(path)
        elif ext == ".docx":
            content = _read_docx(path)
        elif ext == ".csv":
            content = _read_csv_file(path)
        elif ext == ".json":
            content = _read_json_file(path)
        elif ext in _TEXT_EXTENSIONS or ext == "" or size < 1_000_000:
            content = _read_text_file(path)
        else:
            return (
                f"Error: Unsupported file type '{ext}'. "
                "Supported types: PDF, DOCX, CSV, JSON, images, and common text/code files."
            )
    except Exception as exc:
        return f"Error reading file '{path.name}': {exc}"

    return _truncate(header + content)


def list_files(directory_path: str, pattern: str = "*") -> str:
    path = _safe_resolve(directory_path)

    if _is_sensitive_path(path):
        return (
            f"Error: refusing to list sensitive path '{path}'. "
            "SSH/AWS/GnuPG and Yumi's own config directories are blocked from the list_files tool."
        )

    if not path.exists():
        return f"Error: Directory not found: {path}"
    if not path.is_dir():
        return f"Error: '{path}' is not a directory."

    try:
        entries = sorted(path.glob(pattern))
    except Exception as exc:
        return f"Error listing directory: {exc}"

    if not entries:
        return f"No files matching '{pattern}' in {path}"

    lines: list[str] = [f"Contents of {path} (pattern: {pattern}):"]
    dirs = []
    files = []

    for entry in entries:
        if entry.name.startswith("."):
            continue
        if entry.is_dir():
            dirs.append(f"  [DIR]  {entry.name}/")
        elif entry.is_file():
            size = entry.stat().st_size
            files.append(f"  [FILE] {entry.name}  ({_format_size(size)})")

    for d in dirs[:50]:
        lines.append(d)
    for f in files[:200]:
        lines.append(f)

    total = len(dirs) + len(files)
    shown = min(len(dirs), 50) + min(len(files), 200)
    if shown < total:
        lines.append(f"  ... and {total - shown} more entries")

    lines.append(f"\nTotal: {len(dirs)} directories, {len(files)} files")
    return "\n".join(lines)


def _format_size(size: int) -> str:
    n = float(size)
    for unit in ("B", "KB", "MB", "GB"):
        if n < 1024 or unit == "GB":
            return f"{n:.1f} {unit}" if unit != "B" else f"{int(n)} B"
        n /= 1024
    return f"{n:.1f} GB"
